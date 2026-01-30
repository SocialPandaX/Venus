"""
Docx Editor MCP Server
This server provides tools for reading and editing Microsoft Word (.docx) files.
It uses the python-docx library to interact with documents without requiring Word.
"""
import os
import sys
import re
from typing import List, Optional, Iterable, Dict, Any, Union
from docx import Document as DocumentFunction
from docx.document import Document as DocumentClass
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from mcp.server.fastmcp import FastMCP

mcp = FastMCP("Docx Editor")

MCP_PROMPT = """
Docx Editor MCP Server:
- read_docx / read_docx_advanced: 以高保真度读取 Word 文档内容（支持嵌套表格和页眉页脚）。
- create_new_docx / add_heading / add_paragraph / add_list_item: 创建并构建带格式的文档。
- insert_table / set_table_cell_format / merge_table_cells: 专业的表格操作和样式设置。
- replace_text_global / delete_element / insert_content_relative: 精确编辑已有文档内容。
- get_document_info: 获取文档元数据和统计信息。

**重要规范**：
1. 严禁使用普通的 write_file 工具直接创建或写入 .docx 文件，这会导致文件损坏。必须使用本 Server 提供的工具。
2. 针对软著源代码文档（通常内容极长）：
   - 严禁在单次工具调用中传递数十页的代码，这会导致 JSON 截断错误。
   - 正确流程：先调用 create_new_docx 创建文件，然后分多次（例如每 5-10 页代码为一次）调用 add_paragraph 或相关工具追加内容。
   - 每次追加后应确保文件已保存，避免因 Token 限制导致的任务中断。
注意：在容器内操作时，请确保路径以 /workspace/ 或 /template/ 开头。
"""

# Path mapping for Docker container workspace
# This script is in gemini_mcp_client/tools/ or similar
BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
WORKSPACE_HOST = os.path.join(BASE_DIR, "workspace")
WORKSPACE_CONTAINER = "/workspace"

def _resolve_path(path: str) -> str:
    """
    Resolve a container path to a host path if it starts with /workspace.
    Handles both Windows and Linux style paths.
    """
    if not path:
        return path
        
    # Standardize to forward slashes for prefix checking and mapping
    normalized_input = path.replace('\\', '/')
    
    # Check if it starts with the container workspace prefix
    if normalized_input.startswith(WORKSPACE_CONTAINER):
        # Remove the prefix and ensure we don't have a leading slash
        relative_path = normalized_input[len(WORKSPACE_CONTAINER):].lstrip('/')
        # Join with the host workspace path
        resolved = os.path.join(WORKSPACE_HOST, relative_path)
    else:
        # If it's a relative path, assume it's relative to the host workspace
        if not os.path.isabs(path) and not path.startswith(('.', '..')):
             resolved = os.path.join(WORKSPACE_HOST, path)
        else:
             # Absolute host path or relative traversal - return as is (normalized)
             resolved = path
             
    return os.path.normpath(resolved)

def _iter_block_items(parent) -> Iterable:
    """
    Yield each paragraph and table child within *parent*, in document order.
    Each returned value is an instance of either Table or Paragraph.
    """
    if isinstance(parent, DocumentClass):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        raise TypeError("expected Document or _Cell, got %s" % type(parent))

    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)

def _format_block(block, table_level: int = 0) -> List[str]:
    """
    Recursively format a block (Paragraph or Table) into a list of strings.
    'table_level' tracks nesting depth for indentation.
    """
    output = []
    indent = "  " * table_level
    
    if isinstance(block, Paragraph):
        text = block.text.strip()
        if text:
            style = block.style.name.lower()
            if 'heading' in style:
                level = style.replace('heading', '').strip()
                prefix = "#" * (int(level) if level.isdigit() else 1)
                output.append(f"{indent}{prefix} {text}")
            else:
                output.append(f"{indent}{text}")
                
    elif isinstance(block, Table):
        processed_cells = set()
        for row in block.rows:
            row_cells_content = []
            for cell in row.cells:
                cell_id = cell._tc
                if cell_id in processed_cells:
                    # Keep structure for AI but indicate it's a merged area
                    row_cells_content.append("[Merged]")
                    continue
                
                # Recursive call for cell content
                cell_items = []
                for sub_block in _iter_block_items(cell):
                    cell_items.extend(_format_block(sub_block, table_level + 1))
                
                # Format cell text to be cleaner
                cell_text = " ".join(item.strip() for item in cell_items if item.strip())
                row_cells_content.append(cell_text if cell_text else "( )") # Use empty brackets for empty cells
                processed_cells.add(cell_id)
            
            if row_cells_content:
                output.append(f"{indent}| " + " | ".join(row_cells_content) + " |")
                
        if output:
            # Use a more descriptive border
            table_width = len(output[0]) - len(indent)
            border = f"{indent}" + "=" * table_width
            output.insert(0, border)
            output.append(border)
            
    return output

@mcp.tool()
def read_docx(path: str) -> str:
    """
    Alias for read_docx_advanced with default settings.
    """
    return read_docx_advanced(path)

@mcp.tool()
def read_docx_advanced(path: str, include_headers: bool = False, include_footers: bool = False) -> str:
    """
    Read document content with high fidelity, including nested tables.
    Optionally include headers and footers.
    """
    resolved_path = _resolve_path(path)
    if not os.path.exists(resolved_path):
        return f"Error: File not found at {path}"
    
    try:
        doc = DocumentFunction(resolved_path)
        output = []
        
        # Headers
        if include_headers:
            for i, section in enumerate(doc.sections):
                header = section.header
                if header.paragraphs:
                    output.append(f"--- Section {i} Header ---")
                    for p in header.paragraphs:
                        output.extend(_format_block(p))
        
        # Main Body (Sequential & Recursive)
        table_count = 0
        for block in _iter_block_items(doc):
            if isinstance(block, Table):
                table_count += 1
                output.append(f"\n[Table {table_count}]")
            
            output.extend(_format_block(block))
        
        # Footers
        if include_footers:
            for i, section in enumerate(doc.sections):
                footer = section.footer
                if footer.paragraphs:
                    output.append(f"--- Section {i} Footer ---")
                    for p in footer.paragraphs:
                        output.extend(_format_block(p))
                        
        return "\n".join(output) if output else "(Empty Document)"
    except Exception as e:
        return f"Error reading docx: {str(e)}"

@mcp.tool()
def create_new_docx(path: str, title: Optional[str] = None, author: Optional[str] = None) -> str:
    """
    Create a new empty Word document.
    'path' is the container path (e.g. /workspace/new.docx).
    """
    resolved_path = _resolve_path(path)
    try:
        # Ensure directory exists
        os.makedirs(os.path.dirname(resolved_path), exist_ok=True)
        
        doc = DocumentFunction()
        if title:
            doc.core_properties.title = title
            doc.add_heading(title, 0)
        if author:
            doc.core_properties.author = author
            
        doc.save(resolved_path)
        return f"Successfully created document at {path}"
    except Exception as e:
        return f"Error creating docx: {str(e)}"

@mcp.tool()
def add_heading(path: str, text: str, level: int = 1, align: str = "left") -> str:
    """
    Add a heading to the document.
    'level' 0-9. 'align' can be left, center, right.
    """
    resolved_path = _resolve_path(path)
    if not os.path.exists(resolved_path):
        return f"Error: File not found at {path}"
    
    try:
        doc = DocumentFunction(resolved_path)
        h = doc.add_heading(text, level=level)
        
        if align.lower() == "center":
            h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif align.lower() == "right":
            h.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
        doc.save(resolved_path)
        return f"Added heading level {level} to {path}"
    except Exception as e:
        return f"Error adding heading: {str(e)}"

@mcp.tool()
def add_paragraph(path: str, text: str, bold: bool = False, italic: bool = False, 
                  color: Optional[str] = None, size: Optional[int] = None, 
                  align: str = "left") -> str:
    """
    Add a styled paragraph to the document.
    'color' is hex RRGGBB. 'size' is in points.
    """
    resolved_path = _resolve_path(path)
    if not os.path.exists(resolved_path):
        return f"Error: File not found at {path}"
    
    try:
        doc = DocumentFunction(resolved_path)
        p = doc.add_paragraph()
        run = p.add_run(text)
        
        if bold: run.bold = True
        if italic: run.italic = True
        if color:
            run.font.color.rgb = RGBColor.from_string(color)
        if size:
            run.font.size = Pt(size)
            
        if align.lower() == "center":
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif align.lower() == "right":
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
        doc.save(resolved_path)
        return f"Added paragraph to {path}"
    except Exception as e:
        return f"Error adding paragraph: {str(e)}"

@mcp.tool()
def add_list_item(path: str, text: str, style: str = "bulleted") -> str:
    """
    Add a list item (bulleted or numbered).
    'style' can be 'bulleted' or 'numbered'.
    """
    resolved_path = _resolve_path(path)
    if not os.path.exists(resolved_path):
        return f"Error: File not found at {path}"
    
    try:
        doc = DocumentFunction(resolved_path)
        list_style = 'List Bullet' if style.lower() == "bulleted" else 'List Number'
        doc.add_paragraph(text, style=list_style)
        doc.save(resolved_path)
        return f"Added {style} list item to {path}"
    except Exception as e:
        return f"Error adding list item: {str(e)}"

from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

@mcp.tool()
def get_document_info(path: str) -> str:
    """
    Get document metadata and statistics.
    """
    resolved_path = _resolve_path(path)
    if not os.path.exists(resolved_path):
        return f"Error: File not found at {path}"
    
    try:
        doc = DocumentFunction(resolved_path)
        props = doc.core_properties
        
        info = {
            "title": props.title or "N/A",
            "author": props.author or "N/A",
            "created": str(props.created) if props.created else "N/A",
            "modified": str(props.modified) if props.modified else "N/A",
            "paragraphs_count": len(doc.paragraphs),
            "tables_count": len(doc.tables),
            "sections_count": len(doc.sections)
        }
        
        return "\n".join([f"{k}: {v}" for k, v in info.items()])
    except Exception as e:
        return f"Error getting document info: {str(e)}"

@mcp.tool()
def insert_content_relative(path: str, target_text: str, content: str, position: str = "after", 
                             element_type: str = "paragraph") -> str:
    """
    Insert content relative to existing text.
    'position' can be 'before' or 'after'.
    'element_type' can be 'paragraph' or 'heading' (level 1).
    """
    resolved_path = _resolve_path(path)
    if not os.path.exists(resolved_path):
        return f"Error: File not found at {path}"
    
    try:
        doc = DocumentFunction(resolved_path)
        target_idx = -1
        for i, p in enumerate(doc.paragraphs):
            if target_text in p.text:
                target_idx = i
                break
        
        if target_idx == -1:
            return f"Error: Target text '{target_text}' not found"
        
        # Calculate insertion point
        insert_idx = target_idx + 1 if position == "after" else target_idx
        
        # Create new element
        if element_type == "heading":
            new_p = doc.add_heading(content, level=1)
        else:
            new_p = doc.add_paragraph(content)
            
        # Move it to the right place in the XML
        p_element = new_p._element
        doc.paragraphs[insert_idx]._element.addprevious(p_element) if position == "before" else doc.paragraphs[target_idx]._element.addnext(p_element)
        
        doc.save(resolved_path)
        return f"Inserted content {position} '{target_text}'"
    except Exception as e:
        return f"Error inserting relative content: {str(e)}"

@mcp.tool()
def delete_element(path: str, element_type: str, index: int) -> str:
    """
    Delete a paragraph or table by index.
    'element_type' can be 'paragraph' or 'table'.
    """
    resolved_path = _resolve_path(path)
    if not os.path.exists(resolved_path):
        return f"Error: File not found at {path}"
    
    try:
        doc = DocumentFunction(resolved_path)
        if element_type == "paragraph":
            if index >= len(doc.paragraphs):
                return f"Error: Paragraph index {index} out of range"
            p = doc.paragraphs[index]._element
            p.getparent().remove(p)
        elif element_type == "table":
            if index >= len(doc.tables):
                return f"Error: Table index {index} out of range"
            t = doc.tables[index]._element
            t.getparent().remove(t)
        else:
            return "Error: element_type must be 'paragraph' or 'table'"
            
        doc.save(resolved_path)
        return f"Deleted {element_type} at index {index}"
    except Exception as e:
        return f"Error deleting element: {str(e)}"

@mcp.tool()
def replace_text_global(path: str, search_text: str, replace_text: str, include_headers_footers: bool = True) -> str:
    """
    Globally replace text in paragraphs, tables (including nested ones), headers, and footers.
    """
    resolved_path = _resolve_path(path)
    if not os.path.exists(resolved_path):
        return f"Error: File not found at {path}"
    
    try:
        doc = DocumentFunction(resolved_path)
        count = 0
        
        def _process_recursive(parent):
            nonlocal count
            for block in _iter_block_items(parent):
                if isinstance(block, Paragraph):
                    if search_text in block.text:
                        # count actual occurrences in this paragraph
                        occurrences = block.text.count(search_text)
                        block.text = block.text.replace(search_text, replace_text)
                        count += occurrences
                elif isinstance(block, Table):
                    for row in block.rows:
                        for cell in row.cells:
                            _process_recursive(cell)
        
        # 1. Process Main Body
        _process_recursive(doc)
        
        # 2. Process Headers and Footers
        if include_headers_footers:
            for section in doc.sections:
                # Headers
                for header_p in section.header.paragraphs:
                    if search_text in header_p.text:
                        count += header_p.text.count(search_text)
                        header_p.text = header_p.text.replace(search_text, replace_text)
                for header_table in section.header.tables:
                    for row in header_table.rows:
                        for cell in row.cells:
                            _process_recursive(cell)
                            
                # Footers
                for footer_p in section.footer.paragraphs:
                    if search_text in footer_p.text:
                        count += footer_p.text.count(search_text)
                        footer_p.text = footer_p.text.replace(search_text, replace_text)
                for footer_table in section.footer.tables:
                    for row in footer_table.rows:
                        for cell in row.cells:
                            _process_recursive(cell)
                    
        doc.save(resolved_path)
        return f"Replaced {count} occurrences of '{search_text}'"
    except Exception as e:
        return f"Error in global replace: {str(e)}"

@mcp.tool()
def insert_table(path: str, rows: int, cols: int, data: Optional[List[List[str]]] = None, style: str = "Table Grid") -> str:
    """
    Insert a table into the document.
    'data' is an optional 2D list of strings.
    """
    resolved_path = _resolve_path(path)
    if not os.path.exists(resolved_path):
        return f"Error: File not found at {path}"
    
    try:
        doc = DocumentFunction(resolved_path)
        table = doc.add_table(rows=rows, cols=cols)
        table.style = style
        
        if data:
            for r_idx, row_data in enumerate(data):
                if r_idx < rows:
                    for c_idx, cell_value in enumerate(row_data):
                        if c_idx < cols:
                            table.cell(r_idx, c_idx).text = str(cell_value)
                            
        doc.save(resolved_path)
        return f"Inserted {rows}x{cols} table into {path}"
    except Exception as e:
        return f"Error inserting table: {str(e)}"

@mcp.tool()
def set_table_cell_format(path: str, table_index: int, row: int, col: int, 
                          bg_color: Optional[str] = None, bold: bool = False, 
                          align: Optional[str] = None) -> str:
    """
    Format a specific cell in a table.
    'bg_color' is hex RRGGBB. 'align' is left, center, right.
    """
    resolved_path = _resolve_path(path)
    if not os.path.exists(resolved_path):
        return f"Error: File not found at {path}"
    
    try:
        doc = DocumentFunction(resolved_path)
        if table_index >= len(doc.tables):
            return f"Error: Table index {table_index} out of range"
            
        table = doc.tables[table_index]
        cell = table.cell(row, col)
        
        if bg_color:
            shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{bg_color}"/>')
            cell._tc.get_or_add_tcPr().append(shading_elm)
            
        if bold:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.bold = True
                    
        if align:
            alignment = {"left": WD_ALIGN_PARAGRAPH.LEFT, "center": WD_ALIGN_PARAGRAPH.CENTER, "right": WD_ALIGN_PARAGRAPH.RIGHT}.get(align.lower())
            if alignment is not None:
                for p in cell.paragraphs:
                    p.alignment = alignment
                    
        doc.save(resolved_path)
        return f"Formatted cell ({row}, {col}) in table {table_index}"
    except Exception as e:
        return f"Error formatting cell: {str(e)}"

@mcp.tool()
def merge_table_cells(path: str, table_index: int, start_row: int, start_col: int, end_row: int, end_col: int) -> str:
    """
    Merge cells in a table from (start_row, start_col) to (end_row, end_col).
    """
    resolved_path = _resolve_path(path)
    if not os.path.exists(resolved_path):
        return f"Error: File not found at {path}"
    
    try:
        doc = DocumentFunction(resolved_path)
        if table_index >= len(doc.tables):
            return f"Error: Table index {table_index} out of range"
            
        table = doc.tables[table_index]
        cell1 = table.cell(start_row, start_col)
        cell2 = table.cell(end_row, end_col)
        cell1.merge(cell2)
        
        doc.save(resolved_path)
        return f"Merged cells in table {table_index} from ({start_row}, {start_col}) to ({end_row}, {end_col})"
    except Exception as e:
        return f"Error merging cells: {str(e)}"

@mcp.tool()
def add_image_to_docx(path: str, image_path: str, width_inches: float = 6.0) -> str:
    """
    Add an image to an existing .docx file.
    'image_path' is the local path to the image file.
    'width_inches' is the width of the image in the document (default 6.0).
    """
    resolved_path = _resolve_path(path)
    resolved_image_path = _resolve_path(image_path)
    
    if not os.path.exists(resolved_path):
        return f"Error: Document file not found at {path} (Resolved to {resolved_path})"
    if not os.path.exists(resolved_image_path):
        return f"Error: Image file not found at {image_path} (Resolved to {resolved_image_path})"
    
    try:
        doc = DocumentFunction(resolved_path)
        doc.add_picture(resolved_image_path, width=Inches(width_inches))
        doc.save(resolved_path)
        return f"Successfully added image {image_path} to {path}"
    except Exception as e:
        return f"Error adding image to docx: {str(e)}"

if __name__ == "__main__":
    mcp.run()
